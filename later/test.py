import os
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import FastEmbedEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.schema import Document
from loguru import logger


#? PDF to Text Function
def pdf_to_text(pdf_path):
    """Extracts text from a PDF file."""
    text = ""
    with fitz.open(pdf_path) as pdf:
        for page_num in range(len(pdf)):
            page = pdf.load_page(page_num)
            text += page.get_text()
    return text


#? DOCX to Text Function
def docx_to_text(docx_path):
    """Extracts text from a DOCX file."""
    doc = DocxDocument(docx_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text


#? Function to check if the file is a PDF or DOCX
def get_text(file_path):
    """Extracts text from a PDF or DOCX file."""
    _, file_extension = os.path.splitext(file_path)
    if file_extension == ".pdf":
        return pdf_to_text(file_path)
    elif file_extension == ".docx":
        return docx_to_text(file_path)
    else:
        raise ValueError("Unsupported file format.")


#? Function to create document embeddings
def prepare_rag_documents(file_path, chunk_size=1000, chunk_overlap=100):
    """Splits the text from PDF or DOCX and creates a vector store."""
    
    # Extract text based on file type
    text = get_text(file_path)

    # Split text into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", " ", ""]
    )
    docs = text_splitter.create_documents([text])

    # Create document objects with metadata
    formatted_docs = [Document(page_content=doc.page_content, metadata={"source": file_path}) for doc in docs]

    # Embedding model
    embed_model = FastEmbedEmbeddings(model_name="BAAI/bge-base-en-v1.5")
    persist_directory = f"E:/Projects/ragbot/Vector Embeddings/{os.path.basename(file_path)}"
    os.makedirs(persist_directory, exist_ok=True)

    # Store in Chroma vector DB
    vectorstore = Chroma.from_documents(
        documents=formatted_docs,
        embedding=embed_model,
        persist_directory=persist_directory,
        collection_name="rag"
    )

    logger.info(f"Embeddings created successfully in {persist_directory}")
    return vectorstore


if __name__ == "__main__":
    file_path = "E:/Projects/ragbot/documents/hello.docx"
    prepare_rag_documents(file_path)