import os
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import FastEmbedEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.schema import Document
from loguru import logger


class RAGService:
    
    @staticmethod
    def pdf_to_text(pdf_path: str) -> str:
        """Extracts text from a PDF file."""
        text = ""
        with fitz.open(pdf_path) as pdf:
            for page_num in range(len(pdf)):
                page = pdf.load_page(page_num)
                text += page.get_text()
        return text

    @staticmethod
    def docx_to_text(docx_path: str) -> str:
        """Extracts text from a DOCX file."""
        doc = DocxDocument(docx_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text

    @staticmethod
    def get_text(file_path: str) -> str:
        """Extracts text based on the file format."""
        _, file_extension = os.path.splitext(file_path)
        
        if file_extension == ".pdf":
            return RAGService.pdf_to_text(file_path)
        elif file_extension == ".docx":
            return RAGService.docx_to_text(file_path)
        else:
            raise ValueError("Unsupported file format. Only PDF and DOCX are allowed.")

    @staticmethod
    def prepare_rag_documents(file_path: str, chunk_size: int = 1000, chunk_overlap: int = 100):
        """Splits the text and creates embeddings for RAG."""
        
        text = RAGService.get_text(file_path)

        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )
        docs = text_splitter.create_documents([text])

        # Create document objects with metadata
        formatted_docs = [Document(page_content=doc.page_content, metadata={"source": file_path}) for doc in docs]

        # Embedding model
        embed_model = FastEmbedEmbeddings(model_name="BAAI/bge-base-en-v1.5")

        # Create vector store
        persist_directory = f"E:/Projects/ragbot/app/vector_databases/{os.path.basename(file_path)}"
        os.makedirs(persist_directory, exist_ok=True)

        vectorstore = Chroma.from_documents(
            documents=formatted_docs,
            embedding=embed_model,
            persist_directory=persist_directory,
            collection_name="rag"
        )

        logger.info(f"Embeddings stored in {persist_directory}")
        return {"message": "Embeddings created successfully", "path": persist_directory}
